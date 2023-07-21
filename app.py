import os
from docx import Document
from flask import Flask, request, render_template

app = Flask(__name__)


@app.route("/")
def index():
    # HTMLフォームを表示するためのテンプレートを返す
    return render_template("input_form.html")


@app.route("/get_text", methods=["POST"])
def get_text():
    # フォームから入力されたテキストを取得
    doc_number = request.form["doc_number"]
    doc_date = request.form["doc_date"]
    draft_date = request.form["draft_date"]
    drafter = request.form["drafter"]
    summary = request.form["summary"]
    input_text = request.form["input_text"]
    filename = request.form["filename"]

    # Word文書を生成して保存
    generate_docx(filename, input_text, doc_date, doc_number, draft_date, drafter, summary)

    # 応答メッセージを返す
    return f"ファイル '{filename}.docx' が保存されました。"


# 全文を改行及び文字数で区切ってリストに入れる
def slice_txt_into_list(full_txt, slice_length):
    sliced_list = []
    lines = full_txt.split("\n")  # 改行でテキストを分割

    for line in lines:
        while len(line) > slice_length:
            sliced_list.append(line[:slice_length])
            line = line[slice_length:]

        if line:
            sliced_list.append(line)

    return sliced_list


def generate_docx(new_filename, full_txt, date, doc_num, drft_date, drft_person, summary_content):
    # 40字ごとに区切ってリストに入れる
    draft_content = slice_txt_into_list(full_txt, 40)

    # 既存のWord文書を開く
    document = Document("起案文書様式.docx")

    # 文書内の表を取得する
    table = document.tables[0]

    # 文書番号欄のセルを取得
    doc_num_cell = table.cell(1, 8)

    # 文書番号欄の説に文書番号を入力
    doc_num_cell.text = doc_num

    # 文書の日付欄のセルを取得
    date_cell = table.cell(2, 8)

    # 文書の日付欄のセルに日付を入力
    date_cell.text = date

    # 起案日&起案者欄のセルを取得
    draft_date_cell = table.cell(3, 2)

    # 起案日と起案者を同じセルに入力するため連結
    drafter_date = drft_date + "\n"+"\n" + drft_person

    # 起案日&起案者を入力
    draft_date_cell.text = drafter_date

    # 摘要欄のセルを取得
    summary_content_cell = table.cell(4, 4)

    # 摘要欄に内容を入力
    summary_content_title = "摘要\n" + summary_content
    summary_content_cell.text = summary_content_title

    # 本文を起案文書様式の６行目から入力
    line_num = 6

    for draft in draft_content:
        # 表のline_num行目、0列目にあるセルを取得する
        cell = table.cell(line_num, 0)

        for paragraph in cell.paragraphs:
            paragraph.text = paragraph.text.replace("\n", "")
        # セル内のテキストを上書きする
        cell.text = draft
        line_num += 1

    # 上書きした文書を保存、保存場所指定
    output_folder = r"C:\Users\noriko\Desktop\test"
    output_path = os.path.join(output_folder, f"{new_filename}.docx")
    document.save(output_path)


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=8000, debug=True)
